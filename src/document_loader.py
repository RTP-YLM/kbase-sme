"""
Document Loader — E2-2 (PyThaiNLP Section-Aware Chunking)
Replaces text.split() with proper Thai tokenization + section detection.

Output format matches V1 schema (knowledge_sources + knowledge_chunks).
"""
import hashlib
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Chunk parameters (กำหนดใน MOM kick-off: size=512, overlap=128)
DEFAULT_CHUNK_SIZE = 512
DEFAULT_CHUNK_OVERLAP = 128
MIN_SECTION_TOKENS = 100  # section สั้นกว่านี้ → merge กับ section ถัดไป


# ---------------------------------------------------------------------------
# Data classes — ตรงกับ V1 schema
# ---------------------------------------------------------------------------

@dataclass
class SourceInfo:
    """ข้อมูลสำหรับ knowledge_sources row"""
    title: str
    filename: str
    checksum: str
    file_size_bytes: int
    page_count: Optional[int] = None
    doc_type: str = "file"


@dataclass
class Chunk:
    """ข้อมูลสำหรับ knowledge_chunks row (ยังไม่มี embedding — ทำใน E3-1)"""
    content: str
    chunk_index: int
    section_title: Optional[str] = None
    page_number: Optional[int] = None
    token_count: int = 0
    metadata: dict = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Thai Tokenizer
# ---------------------------------------------------------------------------

def _tokenize_thai(text: str) -> list[str]:
    """ตัดคำภาษาไทยด้วย PyThaiNLP newmm engine"""
    try:
        from pythainlp.tokenize import word_tokenize
        return word_tokenize(text, engine="newmm", keep_whitespace=False)
    except ImportError:
        logger.error("pythainlp ไม่ได้ติดตั้ง: pip install pythainlp")
        # fallback: split by whitespace (พัง สำหรับไทย — log error ชัดเจน)
        return text.split()


def _count_tokens(text: str) -> int:
    """นับ token ด้วย PyThaiNLP"""
    return len(_tokenize_thai(text))


# ---------------------------------------------------------------------------
# Section Detection
# ---------------------------------------------------------------------------

@dataclass
class Section:
    title: Optional[str]
    content: str
    page_number: Optional[int] = None


def _detect_sections_text(text: str) -> list[Section]:
    """
    แยก section จาก plain text / markdown
    Header patterns: # หัวข้อ, ## หัวข้อย่อย, หรือบรรทัดสั้น ALLCAPS
    """
    sections: list[Section] = []
    current_title: Optional[str] = None
    current_lines: list[str] = []

    header_re = re.compile(r'^(#{1,3})\s+(.+)$')
    # Thai all-caps heuristic: บรรทัดสั้น (<60 chars) ที่ลงท้ายด้วย \n\n
    short_line_re = re.compile(r'^.{3,60}$')

    lines = text.splitlines()

    def flush():
        content = "\n".join(current_lines).strip()
        if content:
            sections.append(Section(title=current_title, content=content))

    for i, line in enumerate(lines):
        m = header_re.match(line.strip())
        if m:
            flush()
            current_title = m.group(2).strip()
            current_lines = []
        else:
            current_lines.append(line)

    flush()

    # ถ้าไม่มี header เลย → ทั้งหมดเป็น section เดียว
    if not sections:
        sections = [Section(title=None, content=text.strip())]

    return sections


def _detect_sections_docx(file_path: Path) -> list[Section]:
    """แยก section จาก DOCX โดยใช้ Heading styles"""
    from docx import Document  # type: ignore

    doc = Document(file_path)
    sections: list[Section] = []
    current_title: Optional[str] = None
    current_parts: list[str] = []

    def flush():
        content = "\n".join(current_parts).strip()
        if content:
            sections.append(Section(title=current_title, content=content))

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name.startswith("Heading"):
            flush()
            current_title = text
            current_parts = []
        else:
            current_parts.append(text)

    flush()

    if not sections:
        sections = [Section(title=None, content="\n".join(
            p.text for p in doc.paragraphs if p.text.strip()
        ))]

    return sections


def _load_pdf_sections(
    file_path: Path,
    ocr_api_key: Optional[str] = None,
) -> tuple[list[Section], int]:
    """
    โหลด PDF text-based พร้อม page tracking
    ถ้า PDF เป็น scanned (text น้อย) → ใช้ Typhoon OCR อัตโนมัติ
    """
    from pypdf import PdfReader  # type: ignore
    from .ocr import TyphoonOCR, is_scanned_pdf

    # ตรวจสอบก่อนว่าต้องการ OCR หรือไม่
    if is_scanned_pdf(file_path):
        logger.info(f"{file_path.name}: ตรวจพบ scanned PDF → ใช้ Typhoon OCR")
        ocr = TyphoonOCR(api_key=ocr_api_key)
        page_texts, page_count = ocr.pdf_to_texts(file_path)
        sections: list[Section] = []
        for page_num, text in enumerate(page_texts, start=1):
            if text.strip():
                for s in _detect_sections_text(text):
                    sections.append(Section(
                        title=s.title,
                        content=s.content,
                        page_number=page_num,
                    ))
        return sections, page_count

    # PDF มี text layer — อ่านปกติ
    reader = PdfReader(file_path)
    sections = []
    page_count = len(reader.pages)

    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        for s in _detect_sections_text(text):
            sections.append(Section(
                title=s.title,
                content=s.content,
                page_number=page_num,
            ))

    return sections, page_count


def _load_image_sections(
    file_path: Path,
    ocr_api_key: Optional[str] = None,
) -> tuple[list[Section], None]:
    """โหลดรูปภาพ → OCR → sections"""
    from .ocr import TyphoonOCR

    ocr = TyphoonOCR(api_key=ocr_api_key)
    text = ocr.image_to_text(file_path)
    return _detect_sections_text(text), None


def _load_excel_sections(file_path: Path) -> list[Section]:
    """แปลง Excel sheets เป็น markdown tables"""
    import openpyxl  # type: ignore

    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    sections: list[Section] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        # แปลงเป็น markdown table
        def cell(v) -> str:
            return str(v) if v is not None else ""

        header = rows[0]
        md_lines = ["| " + " | ".join(cell(c) for c in header) + " |"]
        md_lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(cell(c) for c in row) + " |")

        sections.append(Section(
            title=sheet_name,
            content="\n".join(md_lines),
        ))

    wb.close()
    return sections


# ---------------------------------------------------------------------------
# Core Chunker
# ---------------------------------------------------------------------------

class ThaiChunker:
    """
    Section-aware chunker สำหรับภาษาไทย
    ใช้ PyThaiNLP tokenize แทน text.split()
    """

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        min_section_tokens: int = MIN_SECTION_TOKENS,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_section_tokens = min_section_tokens

    def chunk_sections(self, sections: list[Section]) -> list[Chunk]:
        """
        แปลง sections → chunks
        - section สั้น (<min_section_tokens) → merge กับ section ถัดไป
        - section ยาว → sliding window ด้วย PyThaiNLP tokens
        """
        # Merge short sections
        merged = self._merge_short_sections(sections)

        chunks: list[Chunk] = []
        global_index = 0

        for section in merged:
            section_chunks = self._chunk_section(section, global_index)
            chunks.extend(section_chunks)
            global_index += len(section_chunks)

        return chunks

    def _merge_short_sections(self, sections: list[Section]) -> list[Section]:
        merged: list[Section] = []
        buffer_content: list[str] = []
        buffer_title: Optional[str] = None
        buffer_page: Optional[int] = None

        def flush_buffer():
            if buffer_content:
                merged.append(Section(
                    title=buffer_title,
                    content="\n".join(buffer_content),
                    page_number=buffer_page,
                ))

        for s in sections:
            token_count = _count_tokens(s.content)
            if token_count < self.min_section_tokens:
                # เพิ่มเข้า buffer
                if not buffer_content:
                    buffer_title = s.title
                    buffer_page = s.page_number
                buffer_content.append(
                    (f"[{s.title}]\n" if s.title and s.title != buffer_title else "") + s.content
                )
            else:
                flush_buffer()
                buffer_content = []
                buffer_title = None
                buffer_page = None
                merged.append(s)

        flush_buffer()
        return merged if merged else sections

    def _chunk_section(self, section: Section, start_index: int) -> list[Chunk]:
        """ตัด section เดียวเป็น chunks ด้วย sliding window บน PyThaiNLP tokens"""
        tokens = _tokenize_thai(section.content)
        total_tokens = len(tokens)

        if total_tokens == 0:
            return []

        if total_tokens <= self.chunk_size:
            return [Chunk(
                content=section.content.strip(),
                chunk_index=start_index,
                section_title=section.title,
                page_number=section.page_number,
                token_count=total_tokens,
                metadata={"doc_type": "section"},
            )]

        chunks: list[Chunk] = []
        step = self.chunk_size - self.chunk_overlap
        idx = start_index

        for i in range(0, total_tokens, step):
            window_tokens = tokens[i: i + self.chunk_size]
            content = "".join(window_tokens).strip()  # Thai ไม่ใส่ space ระหว่างคำ
            if not content:
                continue

            chunks.append(Chunk(
                content=content,
                chunk_index=idx,
                section_title=section.title,
                page_number=section.page_number,
                token_count=len(window_tokens),
                metadata={"doc_type": "section"},
            ))
            idx += 1

            if i + self.chunk_size >= total_tokens:
                break

        return chunks


# ---------------------------------------------------------------------------
# DocumentLoader — public API
# ---------------------------------------------------------------------------

class DocumentLoader:
    """
    โหลดไฟล์ → แยก sections → chunk ด้วย PyThaiNLP
    Returns (SourceInfo, list[Chunk]) สำหรับ upsert เข้า Supabase
    """

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
        ocr_api_key: Optional[str] = None,
    ):
        self.chunker = ThaiChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self._ocr_api_key = ocr_api_key or os.environ.get("TYPHOON_OCR_API_KEY")

    def load_file(self, file_path: str | Path) -> tuple[SourceInfo, list[Chunk]]:
        """
        โหลดและ chunk ไฟล์เดียว

        Returns:
            (SourceInfo, chunks) — SourceInfo สำหรับ knowledge_sources,
                                   chunks สำหรับ knowledge_chunks
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        loaders = {
            ".txt":  self._load_text,
            ".md":   self._load_text,
            ".pdf":  self._load_pdf,
            ".docx": self._load_docx,
            ".xlsx": self._load_excel,
            ".png":  self._load_image,
            ".jpg":  self._load_image,
            ".jpeg": self._load_image,
            ".webp": self._load_image,
        }

        if suffix not in loaders:
            raise ValueError(f"ไม่รองรับ format: {suffix}")

        if not path.exists():
            raise FileNotFoundError(f"ไม่พบไฟล์: {path}")

        checksum = self._sha256(path)
        file_size = path.stat().st_size

        sections, page_count = loaders[suffix](path)
        chunks = self.chunker.chunk_sections(sections)

        # เพิ่ม metadata ต่อ chunk
        for chunk in chunks:
            chunk.metadata.update({
                "source_filename": path.name,
                "doc_type": suffix.lstrip("."),
            })

        source = SourceInfo(
            title=path.stem.replace("_", " ").replace("-", " ").title(),
            filename=path.name,
            checksum=checksum,
            file_size_bytes=file_size,
            page_count=page_count,
            doc_type=suffix.lstrip("."),
        )

        logger.info(
            f"โหลด {path.name}: {len(sections)} sections → {len(chunks)} chunks "
            f"(checksum: {checksum[:8]}...)"
        )
        return source, chunks

    # ------------------------------------------------------------------
    # File-type loaders — คืน (sections, page_count)
    # ------------------------------------------------------------------

    def _load_text(self, path: Path) -> tuple[list[Section], None]:
        text = path.read_text(encoding="utf-8")
        return _detect_sections_text(text), None

    def _load_pdf(self, path: Path) -> tuple[list[Section], int]:
        return _load_pdf_sections(path, ocr_api_key=self._ocr_api_key)

    def _load_image(self, path: Path) -> tuple[list[Section], None]:
        return _load_image_sections(path, ocr_api_key=self._ocr_api_key)

    def _load_docx(self, path: Path) -> tuple[list[Section], None]:
        return _detect_sections_docx(path), None

    def _load_excel(self, path: Path) -> tuple[list[Section], None]:
        return _load_excel_sections(path), None

    # ------------------------------------------------------------------
    # Utilities
    # ------------------------------------------------------------------

    @staticmethod
    def _sha256(path: Path) -> str:
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for block in iter(lambda: f.read(65536), b""):
                h.update(block)
        return h.hexdigest()
